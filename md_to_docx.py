import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'start', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key, value in edge_data.items():
                element.set(qn('w:{}'.format(key)), str(value))

def md_to_docx(md_filename):
    if not os.path.exists(md_filename): return
    doc_filename = md_filename.replace('.md', '.docx')
    doc = Document()
    
    # 1. Page Margin
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    with open(md_filename, 'r') as f:
        lines = f.readlines()

    # State machine for simple parsing
    for line in lines:
        line = line.strip()
        if not line: continue
        
        if line.startswith('# ASGARDIAN'):
            h = doc.add_heading(line.replace('# ', ''), level=0)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in h.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(24)
                run.font.color.rgb = RGBColor(0, 0, 0) # Black for the main title
        
        elif line.startswith('## '):
            h = doc.add_heading(line.replace('## ', ''), level=1)
            for run in h.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(16)
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 73, 125) # Navy Professional
        
        elif line.startswith('### Room:'):
            h = doc.add_heading(line.replace('### ', ''), level=2)
            for run in h.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(14)
                run.font.color.rgb = RGBColor(0, 0, 0)
        
        elif "Device:" in line or "Asset:" in line:
             p = doc.add_paragraph()
             run = p.add_run(line)
             run.bold = True
             run.font.size = Pt(11)

        elif line.startswith('• ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            clean = line[2:]
            parts = clean.split('**')
            for i, part in enumerate(parts):
                r = p.add_run(part)
                if i % 2 != 0: r.bold = True
        else:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                r = p.add_run(part)
                if i % 2 != 0: r.bold = True

    doc.save(doc_filename)
    return doc_filename

if __name__ == "__main__":
    if len(sys.argv) > 1: md_to_docx(sys.argv[1])
